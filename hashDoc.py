#@File: hashDoc.py
#@Author: Diego Martínez Sánchez
#@Brief: The method that has the skeleton for calling tha hash methods, create the word file, take the snapshots, etc

import os, time, hashlib, pyautogui, docx, openpyxl
import MetodosHash, ex
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH#Marca que no existe pero si existe realmente, incluso la doc lo indica de esa manera.

class SacarHash():
    def hash(self, RutaEx):
        noms= MetodosHash.archYrutas()
        xl= ex.Excel()
        doc= docx.Document()

        bandetitulo= 0
        cadeTemp= ""
        dirs2= []
        dirs= []
        DatosFlag= 0
        path= xl.DatosEx(DatosFlag, RutaEx)
        DatosFlag= 1
        extension= xl.DatosEx(DatosFlag, RutaEx)
        fecha_actual= time.strftime("%m/%d/%Y")
        hora_actual= time.strftime("%H:%M:%S")
        cont= 0
        Descripcion= xl.Ex(RutaEx)
        cont_para= 0
        ImaExt= ".png"

        noms.NomCarpe(dirs, path, dirs2)
        dirs2= []
        dirs2= xl.NomCarpesRuta(RutaEx)
        os.system("cls")

        for i in dirs2:
            tupla= ()
            cont_tupla= 0
            a1= 0
            a2= ""
            a3= ""
            bande= 0
            cont_pixeles= 0
            pixeles= 70
            dire= ""
            DatosFlag= 0
            path2= xl.DatosEx(DatosFlag, RutaEx)
            path3= path2
            dire= "\\"+i
            path2+=dire
            path3+=dire
            cadeTemp= noms.NomArch(path2, extension)
            diraux= dirs2
            dirtemporal= diraux[cont].split(" ")
            diraux[cont]= dirtemporal[0]

            print("---------------"+diraux[cont]+"---------------")
            for x in cadeTemp:
                aux= "\\"+x
                path2+= "\\"+x
                hashh, bande= noms.getsha256file(path2)

                if bande== 0:
                    cont_tupla+= 1
                    a1= cont_tupla
                    a2= hashh
                    a3= x
                    tupla2= ((a1, a2, a3),)
                    tupla+= tupla2

                    print("Nombre del archivo: "+x)
                    print("Hash: "+hashh)
                    print("Algoritmo: SHA256")
                    print("")
                    cont_pixeles+= 1
                elif bande>=1:
                    pass
                path2= path2.rstrip(aux)

            print(f"Fecha actual: {fecha_actual}")
            print(f"Hora actual: {hora_actual}")
            print("--------------------------------------------")

            cont_pixeles+= 1
            for h in range(cont_pixeles):
                pixeles+= 50
            pixeles+= 45
            screenshoty= pyautogui.screenshot(region=(1, -0.5, 850, pixeles))
            screenshoty.save(path3+"\\"+diraux[cont]+".png")
            cont+= 1
            os.system("cls")
            pixeles= 70

            try:
                if bandetitulo==0:
                    doc.add_paragraph(xl.TituloWord(RutaEx)).alignment= WD_ALIGN_PARAGRAPH.CENTER
                    bandetitulo= 1
                else:
                    pass
                doc.add_paragraph("\n")
                doc.add_paragraph(Descripcion[cont_para]).alignment= WD_ALIGN_PARAGRAPH.JUSTIFY
            except:
                Descripcion.append("Descripción")
                doc.add_paragraph("\n")
                doc.add_paragraph(Descripcion[cont_para]).alignment= WD_ALIGN_PARAGRAPH.JUSTIFY
            tabla= doc.add_table(rows= 0, cols= 3, style= "Table Grid")
            for c1, c2, c3 in tupla:
                row_cells= tabla.add_row().cells
                row_cells[0].text= str(c1)
                row_cells[1].text= c2
                row_cells[2].text= c3
            cont_para+= 1
            Imas= noms.NomArch(path2, ImaExt)
            for j in range(len(Imas)):
                auxi= "\\"+Imas[j]
                path3+= auxi
                doc.add_picture(path3, width= Inches(6.56))
                path3= path3.rstrip(auxi)

        DatosFlag= 2
        keep= xl.DatosEx(DatosFlag, RutaEx)
        doc.save(keep)